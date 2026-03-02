from fastapi import FastAPI, UploadFile, File, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from sqlalchemy.orm import Session
from apscheduler.schedulers.background import BackgroundScheduler
import io, re, math
from datetime import datetime, timedelta

import PyPDF2
from docx import Document as DocxDocument
from transformers import pipeline

from database import get_db, create_tables, AnalysisSession, ClassificationResult, ReadabilityMetric, Document as DBDocument, SessionLocal

app = FastAPI(title="EduCAS Backend", version="2.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

create_tables()

print("Loading AI model...")
classifier = pipeline("zero-shot-classification", model="typeform/distilbert-base-uncased-mnli")
print("Model ready!")

DOMAIN_LABELS     = ["STEM", "Humanities", "Social Sciences"]
SUBJECT_LABELS    = ["Mathematics", "Physics", "Chemistry", "Biology", "Computer Science",
    "Engineering", "History", "Literature", "Philosophy", "Economics",
    "Political Science", "Sociology", "Psychology", "Geography"]
DIFFICULTY_LABELS = ["Beginner", "Intermediate", "Advanced"]
AUDIENCE_LABELS   = ["Elementary School", "Middle School", "High School", "University"]

# ── AUTO DELETE after 24 hours ─────────────────────────────────────────
def auto_delete_old_records():
    db = SessionLocal()
    try:
        cutoff = datetime.utcnow() - timedelta(hours=24)
        old_sessions = db.query(AnalysisSession).filter(AnalysisSession.created_at < cutoff).all()
        for session in old_sessions:
            db.query(ClassificationResult).filter(ClassificationResult.session_id == session.id).delete()
            db.query(ReadabilityMetric).filter(ReadabilityMetric.session_id == session.id).delete()
            db.delete(session)
        db.commit()
        print(f"🗑️ Auto-deleted {len(old_sessions)} old sessions")
    except Exception as e:
        print(f"Auto-delete error: {e}")
    finally:
        db.close()

scheduler = BackgroundScheduler()
scheduler.add_job(auto_delete_old_records, 'interval', hours=1)
scheduler.start()

def extract_text_from_pdf(content: bytes) -> str:
    reader = PyPDF2.PdfReader(io.BytesIO(content))
    return " ".join(page.extract_text() or "" for page in reader.pages)

def extract_text_from_docx(content: bytes) -> str:
    doc = DocxDocument(io.BytesIO(content))
    return " ".join(para.text for para in doc.paragraphs)

def count_syllables(word: str) -> int:
    word = word.lower().strip(".,!?;:")
    if len(word) <= 3: return 1
    vowels = "aeiouy"
    count, prev_vowel = 0, False
    for char in word:
        is_vowel = char in vowels
        if is_vowel and not prev_vowel: count += 1
        prev_vowel = is_vowel
    if word.endswith("e"): count -= 1
    return max(1, count)

def calculate_readability(text: str) -> dict:
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 5]
    words = re.findall(r'\b[a-zA-Z]+\b', text)
    if not sentences or not words:
        return {"fleschScore": None, "smogScore": None, "readingLevel": None, "avgSentenceLength": None, "complexSentences": None}

    num_sentences = len(sentences)
    num_words = len(words)
    syllable_counts = [count_syllables(w) for w in words]
    num_syllables = sum(syllable_counts)

    avg_sentence_len = round(num_words / num_sentences, 1)
    avg_syllables = num_syllables / num_words

    # Flesch Reading Ease
    flesch = round(206.835 - (1.015 * avg_sentence_len) - (84.6 * avg_syllables), 1)
    flesch = max(0, min(100, flesch))

    # SMOG Grade
    polysyllabic = sum(1 for c in syllable_counts if c >= 3)
    if num_sentences >= 30:
        smog = round(3 + math.sqrt(polysyllabic * (30 / num_sentences)), 1)
    else:
        smog = round(3 + math.sqrt(polysyllabic), 1)

    complex_count = sum(1 for s in sentences if len(s.split()) > 25)
    fk_grade = round(0.39 * avg_sentence_len + 11.8 * avg_syllables - 15.59, 1)
    fk_grade = max(0, fk_grade)

    if fk_grade <= 5: level = "Elementary"
    elif fk_grade <= 8: level = "Middle School"
    elif fk_grade <= 12: level = "High School"
    else: level = "University"

    return {"fleschScore": flesch, "smogScore": smog, "readingLevel": level,
            "avgSentenceLength": avg_sentence_len, "complexSentences": complex_count}

def classify_text(text: str) -> dict:
    snippet = text[:1000]
    domain_result     = classifier(snippet, DOMAIN_LABELS)
    subject_result    = classifier(snippet, SUBJECT_LABELS)
    difficulty_result = classifier(snippet, DIFFICULTY_LABELS)
    audience_result   = classifier(snippet, AUDIENCE_LABELS)
    domain       = domain_result["labels"][0]
    confidence   = round(domain_result["scores"][0], 4)
    subject      = subject_result["labels"][0]
    difficulty   = difficulty_result["labels"][0]
    suitable_for = audience_result["labels"][0]
    sentences = re.split(r'[.!?]+', text[:3000])
    sentences = [s.strip() for s in sentences if len(s.split()) > 5]
    summary = ". ".join(sentences[:3]) + "." if sentences else text[:300]
    recommendations = (
        f"This {difficulty.lower()}-level {subject} content is best suited for {suitable_for} students. "
        f"Educators can use this material to introduce or reinforce key concepts in {domain}. "
        f"Consider supplementing with visual aids or hands-on activities for better engagement."
    )
    return {"domain": domain, "subject": subject, "difficulty": difficulty,
            "suitableFor": suitable_for, "confidence": confidence,
            "summary": summary, "recommendations": recommendations}

def save_to_db(db: Session, classification: dict, readability: dict, filename: str = None, input_type: str = "text"):
    try:
        doc_id = None
        if filename:
            ext = filename.split(".")[-1] if "." in filename else "unknown"
            db_doc = DBDocument(filename=filename, file_type=ext)
            db.add(db_doc)
            db.flush()
            doc_id = db_doc.id
        session = AnalysisSession(document_id=doc_id, input_type=input_type)
        db.add(session)
        db.flush()
        clf = ClassificationResult(
            session_id=session.id, domain=classification["domain"],
            subject=classification["subject"], difficulty=classification["difficulty"],
            suitable_for=classification["suitableFor"], confidence=classification["confidence"],
            summary=classification["summary"], recommendations=classification["recommendations"]
        )
        db.add(clf)
        rdbl = ReadabilityMetric(
            session_id=session.id, flesch_score=readability.get("fleschScore"),
            reading_level=readability.get("readingLevel"),
            avg_sentence_length=readability.get("avgSentenceLength"),
            complex_sentences=readability.get("complexSentences")
        )
        db.add(rdbl)
        db.commit()
        print(f"✅ Saved to DB - session {session.id}")
    except Exception as e:
        db.rollback()
        print(f"⚠️ DB save failed: {e}")

class TextInput(BaseModel):
    text: str

@app.get("/")
def root():
    return {"message": "EduCAS FastAPI Backend is running!", "version": "2.0.0"}

@app.post("/analyze")
def analyze_text(body: TextInput, db: Session = Depends(get_db)):
    if not body.text.strip():
        raise HTTPException(status_code=400, detail="No text provided")
    try:
        classification = classify_text(body.text)
        readability = calculate_readability(body.text)
        save_to_db(db, classification, readability, input_type="text")
        return {**classification, "readability": readability}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload")
async def upload_file(file: UploadFile = File(...), db: Session = Depends(get_db)):
    content = await file.read()
    if len(content) > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File exceeds 5MB limit")
    filename = file.filename.lower()
    try:
        if filename.endswith(".pdf"):
            text = extract_text_from_pdf(content)
        elif filename.endswith(".docx"):
            text = extract_text_from_docx(content)
        elif filename.endswith(".txt"):
            text = content.decode("utf-8")
        else:
            raise HTTPException(status_code=400, detail="Use PDF, DOCX, or TXT")
        if not text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from file")
        classification = classify_text(text)
        readability = calculate_readability(text)
        save_to_db(db, classification, readability, filename=file.filename, input_type="file")
        return {**classification, "readability": readability}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))